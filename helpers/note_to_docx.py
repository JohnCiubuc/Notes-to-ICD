#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Aug  5 21:06:31 2022

@author: John Ciubuc
"""
import math
import json
from docx import Document 
from docx.shared import RGBColor

def noteSection(note,item, subItem = -1):
    if subItem == -1:
        for att in (x[5] for x in note if x[3] == item):
            return att.replace('\x00', 'ti') # Fixes pdf pasting
    else:
        for att in (x[5] for x in note if x[3] == item and x[4] == subItem):
            return att.replace('\x00', 'ti') # Fixes pdf pasting
    return ""


class lazy_document():
    _text = ''
        
    def add_heading(self, text, ignore):
        self._text = self._text + text + '\r\n\r\n'
    def add_paragraph(self, text):
        self.add_heading(text, 0)
    def save(self, filename):
        filename=filename.replace('docx', 'txt')
        with open(filename, 'w') as f:
            f.write(self._text)
            f.close()

def exportStudentNotes_text(note_content, id):
    note = note_content
    
    document = lazy_document()
    document.add_heading(f'Student Note ID: {id}', 0)
    document.add_heading('Chief Complaint', 1)
    document.add_paragraph(noteSection(note,20))
    document.add_heading('History of Presenting Illness', 1)
    document.add_paragraph(noteSection(note,21))
    document.add_heading('Review of Systems', 1)
    document.add_paragraph(noteSection(note,22))
    document.add_heading('History', 1)
    document.add_heading('Past Medical History', 2)
    document.add_paragraph(noteSection(note,23))
    document.add_heading('Past Surgical History', 2)
    document.add_paragraph(noteSection(note,24))
    document.add_heading('Medications', 2)
    document.add_paragraph(noteSection(note,25))
    document.add_heading('Allergies', 2)
    document.add_paragraph(noteSection(note,26))
    document.add_heading('Family History', 2)
    document.add_paragraph(noteSection(note,27))
    document.add_heading('Social History', 2)
    document.add_paragraph(noteSection(note,28))
    document.add_heading('Physical Exam', 1)
    document.add_heading('Vitals', 2)
    
    doc = f"""Heart Rate: {noteSection(note,29,0)}, Blood Pressure: {noteSection(note,29,1)}
Respiratory Rate: {noteSection(note,29,2)},  O2 Sat: {noteSection(note,29,3)}
Weight: {noteSection(note,29,4)}, Height: {noteSection(note,29,5)}"""
    
    document.add_paragraph(doc)
    document.add_heading('Exam', 2)
    document.add_paragraph(noteSection(note,39))
    document.add_heading('Data', 1)
    document.add_paragraph(noteSection(note,32))
    document.add_heading('Assessment and Plan', 1)
    document.add_heading('Summary Statement', 2)
    
    doc = f"""This is a {noteSection(note,38,1)} year old {noteSection(note,38,3)}, who is presenting today for {noteSection(note,38,6)}
The patient has a pertinent history of {noteSection(note,38,9)}
Patient's exam is remarkable for {noteSection(note,38,12)}
Patient's data is remarkable for {noteSection(note,38,15)}"""

    document.add_paragraph(doc)
    
    document.add_heading('Impression', 2)
    document.add_heading('Primary Diagnosis:', 3)
    document.add_paragraph(f"{noteSection(note, 35,0)}")
    document.add_heading('Differential Diagnosis:', 3)
    document.add_paragraph(f"{noteSection(note, 35,1)}")
    document.add_heading('Severity Assessment:', 3)
    document.add_paragraph(f"{noteSection(note, 35,2)}")
    document.add_heading('Clinical Trajectory and Contingency:', 3)
    document.add_paragraph(f"{noteSection(note, 35,3)}")
    
    try:
        for problem_id in range(0,math.ceil(max(x[4] for x in note_content if x[3] == 36)/4)):
                document.add_heading(f'Problem {problem_id+1}:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+0))
                document.add_heading('Differential DX:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+1))
                document.add_heading('Diagnostic Plan:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+2))
                document.add_heading('Treatment Plan:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+3))
    except:
        print('no problem list')
                

    # document.add_paragraph(doc)

    document.save(f'Student Note {id}.docx')

def exportStudentNotes_docx(note_content, id):
    note = note_content
    # Add a heading of level 0 (Also called Title)
    document = Document()
   # Choosing the top most section of the page
    section = document.sections[0]
     
    # Selecting the header
    header = section.header
     
    # Selecting the paragraph already present in
    # the header section
    header_para = header.paragraphs[0]
    
    header_para = header_para.add_run("\tThis document was exported from MAENI and may contain sensitive medical information.")
     
    # Adding the centred zoned header
    # header_para.text = 
    font = header_para.font
    font.color.rgb = RGBColor.from_string('8e0000')
    document.add_heading(f'Student Note ID: {id}', 0)
    document.add_heading('Chief Complaint', 1)
    document.add_paragraph(noteSection(note,20))
    document.add_heading('History of Presenting Illness', 1)
    document.add_paragraph(noteSection(note,21))
    document.add_heading('Review of Systems', 1)
    document.add_paragraph(noteSection(note,22))
    document.add_heading('History', 1)
    document.add_heading('Past Medical History', 2)
    document.add_paragraph(noteSection(note,23))
    document.add_heading('Past Surgical History', 2)
    document.add_paragraph(noteSection(note,24))
    document.add_heading('Medications', 2)
    document.add_paragraph(noteSection(note,25))
    document.add_heading('Allergies', 2)
    document.add_paragraph(noteSection(note,26))
    document.add_heading('Family History', 2)
    document.add_paragraph(noteSection(note,27))
    document.add_heading('Social History', 2)
    document.add_paragraph(noteSection(note,28))
    document.add_heading('Physical Exam', 1)
    document.add_heading('Vitals', 2)
    
    doc = f"""Heart Rate: {noteSection(note,29,0)}, Blood Pressure: {noteSection(note,29,1)}
Respiratory Rate: {noteSection(note,29,2)},  O2 Sat: {noteSection(note,29,3)}
Weight: {noteSection(note,29,4)}, Height: {noteSection(note,29,5)}"""
    
    document.add_paragraph(doc)
    document.add_heading('Exam', 2)
    document.add_paragraph(noteSection(note,39))
    document.add_heading('Data', 1)
    document.add_paragraph(noteSection(note,32))
    document.add_heading('Assessment and Plan', 1)
    document.add_heading('Summary Statement', 2)
    
    doc = f"""This is a {noteSection(note,38,1)} year old {noteSection(note,38,3)}, who is presenting today for {noteSection(note,38,6)}
The patient has a pertinent history of {noteSection(note,38,9)}
Patient's exam is remarkable for {noteSection(note,38,12)}
Patient's data is remarkable for {noteSection(note,38,15)}"""

    document.add_paragraph(doc)
    
    document.add_heading('Impression', 2)
    document.add_heading('Primary Diagnosis:', 3)
    document.add_paragraph(f"{noteSection(note, 35,0)}")
    document.add_heading('Differential Diagnosis:', 3)
    document.add_paragraph(f"{noteSection(note, 35,1)}")
    document.add_heading('Severity Assessment:', 3)
    document.add_paragraph(f"{noteSection(note, 35,2)}")
    document.add_heading('Clinical Trajectory and Contingency:', 3)
    document.add_paragraph(f"{noteSection(note, 35,3)}")
    
    try:
        for problem_id in range(0,math.ceil(max(x[4] for x in note_content if x[3] == 36)/4)):
                document.add_heading(f'Problem {problem_id+1}:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+0))
                document.add_heading('Differential DX:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+1))
                document.add_heading('Diagnostic Plan:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+2))
                document.add_heading('Treatment Plan:', 3)
                document.add_paragraph(noteSection(note_content,36,problem_id*4+3))
    except:
        print('no problem list')
                

    # document.add_paragraph(doc)

    document.save(f'Student Note {id}.docx')